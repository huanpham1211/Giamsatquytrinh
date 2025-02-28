import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.oxml.ns import qn  # Import the namespace
from io import BytesIO
import datetime
import re

st.set_page_config(page_title="Báo cáo giám sát", page_icon="👨‍⚕️")


def normalize_text(text):
    if isinstance(text, str):
        return text.strip().lower().replace("đạt", "đạt").replace("có", "có")
    return text

def process_excel(file):
    df = pd.read_excel(file)
    df = df.applymap(normalize_text)

    # Identify relevant columns
    standard_columns = df.columns[:11].tolist()
    step_columns = [col for col in df.columns if col not in standard_columns][:-1]
    
    # Extract rows where "Nhận xét" is not "Đạt"
    non_dạt_comments = df[df["Nhận xét:"].str.lower() != "đạt"][["Điều dưỡng thực hiện", "Khoa đánh giá", "Nhận xét:"]]
    
    # Calculate "Đạt" (from step columns) and "Có" (from compliance columns)
    compliance_percent = df.groupby("Khoa đánh giá")[df.columns[5:11]].apply(lambda x: (x == "có").mean() * 100)
    step_percent = df.groupby("Khoa đánh giá")[step_columns].apply(lambda x: (x == "đạt").mean() * 100)

    # Compute the overall mean percentage for "Có" and "Đạt" per department
    dept_report = pd.concat([compliance_percent.mean(axis=1), step_percent.mean(axis=1)], axis=1)
    dept_report.columns = ["Xác thực người bệnh(%)", "Bảng kiểm(%)"]  # Ensure exactly two columns

    # Step summary (Percentage of "Đạt")
    step_summary = df[step_columns].apply(lambda x: (x == "đạt").mean() * 100).round(1)

    # Success distribution
    df["SuccessRate"] = df[step_columns].apply(lambda row: (row == "đạt").mean() * 100, axis=1).round(1)
    success_distribution = pd.Series({
        "> 90% Tốt": (df["SuccessRate"] > 90).mean() * 100,
        "70-90% Khá": ((df["SuccessRate"] > 70) & (df["SuccessRate"] <= 90)).mean() * 100,
        "50-70% Trung bình": ((df["SuccessRate"] > 50) & (df["SuccessRate"] <= 70)).mean() * 100,
        "< 50% Kém": (df["SuccessRate"] <= 50).mean() * 100
    })

    # Identify top 5 mistakes
    total_records = len(df)
    compliance_mistakes_percentage = df.iloc[:, 5:11].apply(lambda x: (x != "có").mean() * 100)
    step_mistakes_percentage = df[step_columns].apply(lambda x: (x != "đạt").mean() * 100)

    # Combine both mistake percentages and select top 5
    mistake_percentages = pd.concat([compliance_mistakes_percentage, step_mistakes_percentage])
    top_5_mistakes_percentage = mistake_percentages.nlargest(5)

    return step_summary, success_distribution, dept_report, top_5_mistakes_percentage, total_records


def format_header_text(cell, text):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 1  # Center align text
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13)  # Adjust size if needed

def add_heading(doc, text):
    paragraph = doc.add_paragraph("\n" + text)
    run = paragraph.runs[0]
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
    run.font.size = Pt(13)  # Adjust heading size if needed
    paragraph.alignment = 0  # Align to left (default)

def generate_word_report_with_charts(step_summary, success_distribution, dept_report, top_5_mistakes_percentage, total_records, non_dạt_comments):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # Header
    header_table = doc.add_table(rows=1, cols=2)
    hdr_cells = header_table.rows[0].cells
    format_header_text(hdr_cells[0], "BỆNH VIỆN HÙNG VƯƠNG\nPHÒNG ĐIỀU DƯỠNG\nSố:    /BC-PĐD")
    format_header_text(hdr_cells[1], "Cộng hòa xã hội chủ nghĩa Việt Nam\nĐộc lập - Tự do - Hạnh Phúc")

    add_heading(doc, "I. THÔNG TIN CHUNG:")
    doc.add_paragraph("- Thời gian giám sát: ____________________________")
    doc.add_paragraph("- Số lượt hồ sơ kiểm tra: ________________________")
    doc.add_paragraph("- Số lượt giám sát quy trình: ____________________")
    add_heading(doc, "II. KẾT QUẢ GIÁM SÁT:")

    # **1st Chart: Tỷ lệ đạt từng bước (Horizontal Bar Chart with Correct Step Extraction)**
    fig, ax = plt.subplots(figsize=(6, 4))
    colors = sns.color_palette("husl", len(step_summary))
    bars = step_summary.sort_values().plot(kind="barh", ax=ax, color=colors)
    
    # Extract dynamic labels (Bước X)
    extracted_labels = [re.search(r"Bước \d+", col) for col in step_summary.index]
    extracted_labels = [match.group(0) if match else col for match in extracted_labels]
    
    ax.set_yticks(range(len(step_summary)))
    ax.set_yticklabels(extracted_labels, fontsize=10)
    ax.set_xlabel("Tỷ lệ (%)")
    ax.set_ylabel("Bước")
    ax.set_title("Biểu đồ: Tỷ lệ đạt từng bước")
    
    # Remove legend & add labels inside bars
    ax.legend().remove()
    for bar, label in zip(bars.patches, extracted_labels):
        width = bar.get_width()
        ax.text(width + 1, bar.get_y() + bar.get_height()/2, f"{width:.1f}%", 
                ha="left", va="center", fontsize=10, color="black")
    
    img_chart1 = BytesIO()
    plt.savefig(img_chart1, format="png", bbox_inches="tight")
    img_chart1.seek(0)
    doc.add_paragraph("\nBiểu đồ: Tỷ lệ đạt từng bước\n", style='Heading 3')
    doc.add_picture(img_chart1, width=Inches(6))
    
    # **1st Table: Percentage of Bước đạt**
    doc.add_paragraph("\nBảng: Percentage of Bước đạt\n", style='Heading 3')
    table = doc.add_table(rows=len(step_summary) + 1, cols=2)
    table.cell(0, 0).text = "Bước"
    table.cell(0, 1).text = "Tỷ lệ (%)"
    
    for key, value in zip(extracted_labels, step_summary.sort_values().values):
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = f"{value:.1f}%"



    # **2nd Chart: Phân bố nhân viên đạt tiêu chuẩn**
    fig, ax = plt.subplots(figsize=(6, 4))
    success_distribution.plot.pie(autopct="%1.1f%%", ax=ax, startangle=90, colors=sns.color_palette("pastel"))
    ax.set_ylabel("")
    ax.set_title("Biểu đồ: Phân bố nhân viên đạt tiêu chuẩn")

    img_chart2 = BytesIO()
    plt.savefig(img_chart2, format="png", bbox_inches="tight")
    img_chart2.seek(0)
    doc.add_paragraph("\nBiểu đồ: Phân bố nhân viên đạt tiêu chuẩn\n", style='Heading 3')
    doc.add_picture(img_chart2, width=Inches(6))

    # **2nd Table: Percentage of Success Rate**
    doc.add_paragraph("\nBảng: Phần trăm nhân viên đạt tiêu chuẩn\n", style='Heading 3')
    table = doc.add_table(rows=len(success_distribution) + 1, cols=2)
    table.cell(0, 0).text = "Hạng mục"
    table.cell(0, 1).text = "Tỷ lệ (%)"

    for key, value in success_distribution.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = f"{value:.1f}%"

    # **3rd Chart: Tỷ lệ đạt theo khoa (With percentage labels on bars)**
    fig, ax = plt.subplots(figsize=(8, 5))
    dept_report.plot(kind='bar', ax=ax, color=['#1f77b4', '#ff7f0e'])
    
    ax.set_title("Biểu đồ: Tỷ lệ đạt theo khoa")
    ax.set_xlabel("Khoa đánh giá")
    ax.set_ylabel("Tỷ lệ (%)")
    ax.set_xticklabels(dept_report.index, rotation=45, ha="right")
    
    # Add percentage labels on bars
    for bar in ax.patches:
        height = bar.get_height()
        if height > 0:
            ax.text(bar.get_x() + bar.get_width()/2, height + 1, f"{height:.1f}%", 
                    ha="center", fontsize=10, color="black")
    
    img_chart3 = BytesIO()
    plt.savefig(img_chart3, format="png", bbox_inches="tight")
    img_chart3.seek(0)
    doc.add_paragraph("\nBiểu đồ: Tỷ lệ đạt theo khoa\n", style='Heading 3')
    doc.add_picture(img_chart3, width=Inches(6))

    # **4th Chart: Top 5 Common Mistakes**
    fig, ax = plt.subplots(figsize=(6, 4))
    sns.barplot(y=top_5_mistakes_percentage.index, x=top_5_mistakes_percentage.values, ax=ax, color="red")

    ax.set_xlabel("Tỷ lệ (%)")
    ax.set_title("Biểu đồ: Top 5 Sai sót phổ biến nhất")

    # Add percentage labels on bars
    for bar in ax.patches:
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, f"{bar.get_width():.1f}%", 
                ha="left", va="center", fontsize=10, color="black")

    # Save the chart
    img_chart4 = BytesIO()
    plt.savefig(img_chart4, format="png", bbox_inches="tight")
    img_chart4.seek(0)

    # Add chart to Word document
    doc.add_paragraph("\nBiểu đồ: Top 5 Sai sót phổ biến nhất\n", style='Heading 3')
    doc.add_picture(img_chart4, width=Inches(6))

    # **4th Table: Top 5 Common Mistakes**
    doc.add_paragraph("\nBảng: Top 5 Sai sót phổ biến nhất\n", style='Heading 3')
    table = doc.add_table(rows=len(top_5_mistakes_percentage) + 1, cols=2)
    table.cell(0, 0).text = "Hạng mục"
    table.cell(0, 1).text = "Tỷ lệ (%)"

    for key, value in top_5_mistakes_percentage.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = f"{value:.1f}%"

    # **Finalization**
    # Add formatted headings
    add_heading(doc, "III. ĐÁNH GIÁ VÀ NHẬN XÉT:")
    doc.add_paragraph("- Điểm mạnh:\n\n- Điểm cần cải thiện:\n")

    # Add extracted non-Đạt comments
    doc.add_paragraph("\nCác trường hợp chưa đạt:")
    for index, row in non_dạt_comments.iterrows():
        doc.add_paragraph(f"- {row['Điều dưỡng thực hiện']} ({row['Khoa đánh giá']}): {row['Nhận xét:']}", style='List Bullet')

    add_heading(doc, "IV. KẾT LUẬN:")
    doc.add_paragraph("- Tóm tắt tình hình:\n")

    add_heading(doc, "V. ĐỀ XUẤT:")
    doc.add_paragraph("- Nhắc nhở nhân viên y tế:\n- Kiểm tra định kỳ:\n- Phối hợp các khoa:\n")

    add_heading(doc, "VI. PHƯƠNG HƯỚNG:")
    doc.add_paragraph("Tiếp tục kiểm tra, giám sát hồ sơ bệnh án, phát hiện các thiếu sót...")
    
    
    # **Finalization (Right-Aligned Date & "Người báo cáo")**


    # **Finalization (Right-Aligned, Two-Column Table)**
    final_table = doc.add_table(rows=1, cols=2)

    # Set column widths: First column (2x width), Second column (1x width)
    final_table.columns[0].width = Inches(4)
    final_table.columns[1].width = Inches(2)

    # First column (empty space)
    final_table.rows[0].cells[0].text = ""

    # Second column (Date + "Người báo cáo")
    final_cell = final_table.rows[0].cells[1]
    paragraph = final_cell.paragraphs[0]
    paragraph.alignment = 1  # Center text inside the second column

    # Add formatted text
    paragraph.add_run("\nNgày {} tháng {} năm {}\n".format(datetime.datetime.now().day, 
                                                           datetime.datetime.now().month, 
                                                           datetime.datetime.now().year)).bold = True
    paragraph.add_run("Người báo cáo").bold = True

    # Center vertically (adjust cell properties)
    # Center vertically (adjust cell properties)
    tbl = final_table._element
    tbl.set("alignment", "right")  # Ensure the whole table is right-aligned

    for cell in final_table.columns[1].cells:
        tc = cell._element
        tcPr = OxmlElement('w:tcPr')
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), "center")  # Use qn() to ensure namespace is correct
        tcPr.append(vAlign)
        tc.append(tcPr)



    # Save document to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ---- Streamlit UI ----
st.title("📊 Báo cáo tóm tắt Giám sát điều dưỡng & Thực hiện Quy trình")

# File uploader
uploaded_file = st.file_uploader("📂 Tải lên file Excel", type=["xlsx"])

if uploaded_file:
    step_summary, success_distribution, dept_report, top_5_mistakes, total_records = process_excel(uploaded_file)  # Unpacking 5 values

    # Generate Word report
    word_buffer = generate_word_report_with_charts(step_summary, success_distribution, dept_report, top_5_mistakes, total_records)

    # Provide download button
    st.download_button(
        label="📥 Tải xuống báo cáo",
        data=word_buffer,
        file_name="BaoCao_GiamSatDieuDuong.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.success("✅ Báo cáo đã được tạo thành công!")

